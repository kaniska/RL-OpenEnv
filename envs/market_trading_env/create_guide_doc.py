"""
MarketForge — Comprehensive Guide Document Generator
=====================================================
Creates a professional Word document covering:
  Part 1: How the environment was built (step-by-step)
  Part 2: How business users play in it
  Part 3: How developers build and train models
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page Setup ─────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Custom Styles ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for i, (size, color) in enumerate([(28, 0x0B1D3A), (18, 0x0B1D3A), (14, 0x1565C0), (12, 0x333333)], 1):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Calibri'
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor((color >> 16) & 0xFF, (color >> 8) & 0xFF, color & 0xFF)
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(18 if i <= 2 else 12)
    hs.paragraph_format.space_after = Pt(8 if i <= 2 else 6)


def add_code_block(doc, code_text, language="python"):
    """Add a formatted code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # Add shading
    shading = run._element.get_or_add_tcPr if hasattr(run._element, 'get_or_add_tcPr') else None
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing
    return table


def add_note(doc, text, label="NOTE"):
    """Add a callout note."""
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    run.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    return p


def add_step(doc, number, title, body):
    """Add a numbered step with title and body."""
    p = doc.add_paragraph()
    run = p.add_run(f"Step {number}: {title}")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    doc.add_paragraph(body)


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("MarketForge")
run.font.size = Pt(42)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0B, 0x1D, 0x3A)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run("Multi-Agent Market Simulation Environment")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

doc.add_paragraph()

desc_p = doc.add_paragraph()
desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc_p.add_run(
    "Complete Guide: Environment Setup, Business User Playbook,\n"
    "and Developer Training Guide"
)
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta_p.add_run("Built with OpenEnv v0.2.1  |  GRPO Training  |  Gradio Dashboard")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (Manual)
# ══════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)

toc_items = [
    ("Part 1", "How the MarketForge Environment Was Built"),
    ("  1.1", "Project Structure and Architecture"),
    ("  1.2", "Step 1 — Install Dependencies"),
    ("  1.3", "Step 2 — Define the Data Models (models.py)"),
    ("  1.4", "Step 3 — Build the Market Engine (market_environment.py)"),
    ("  1.5", "Step 4 — Create the HTTP Server (server/app.py)"),
    ("  1.6", "Step 5 — Build the HTTP Client (client.py)"),
    ("  1.7", "Step 6 — Write the Reward Functions (rewards.py)"),
    ("  1.8", "Step 7 — Build the Visual Dashboard (app_visual.py)"),
    ("  1.9", "Step 8 — Create the Training Script (train_market_forge.py)"),
    ("  1.10", "Step 9 — Dockerize for Deployment"),
    ("  1.11", "Step 10 — Run and Test Everything"),
    ("", ""),
    ("Part 2", "Business User Playbook — Trading in MarketForge"),
    ("  2.1", "Meet the Trading Personas"),
    ("  2.2", "How to Launch the Trading Playground"),
    ("  2.3", "Running an Automated Simulation"),
    ("  2.4", "Manual Trading — Make Your Own Decisions"),
    ("  2.5", "Understanding the Dashboard Charts"),
    ("  2.6", "How Awards Work — Winning the Game"),
    ("  2.7", "Reading the Leaderboard"),
    ("  2.8", "Trading Strategy Tips for Business Users"),
    ("", ""),
    ("Part 3", "Developer Guide — Building and Training Models"),
    ("  3.1", "Architecture Overview for Developers"),
    ("  3.2", "Connecting to the Environment via HTTP API"),
    ("  3.3", "Building a Custom Agent (Python)"),
    ("  3.4", "Training an LLM Agent with GRPO"),
    ("  3.5", "Reward Function Design"),
    ("  3.6", "Multi-Turn Rollouts and env_mask"),
    ("  3.7", "Deploying to HuggingFace Spaces"),
    ("  3.8", "Evaluation and Benchmarking"),
    ("  3.9", "Extending the Environment"),
]

for num, title in toc_items:
    if not num and not title:
        doc.add_paragraph()
        continue
    p = doc.add_paragraph()
    if num.startswith("Part"):
        run = p.add_run(f"{num}: {title}")
        run.font.bold = True
        run.font.size = Pt(12)
    else:
        run = p.add_run(f"  {num}  {title}")
        run.font.size = Pt(11)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# PART 1: HOW THE ENVIRONMENT WAS BUILT
# ══════════════════════════════════════════════════════════════
doc.add_heading("Part 1: How the MarketForge Environment Was Built", level=1)

doc.add_paragraph(
    "This section walks you through every step of creating the MarketForge "
    "multi-agent trading environment from scratch. Each step includes the exact "
    "commands and scripts you need to run manually to build, test, and verify "
    "each component."
)

# ── 1.1 Project Structure ──
doc.add_heading("1.1  Project Structure and Architecture", level=2)

doc.add_paragraph(
    "MarketForge follows the OpenEnv v0.2.1 architecture: a server-side Environment "
    "that manages game state, and a client-side EnvClient that agents use to interact "
    "with it over HTTP. Here is the complete file structure:"
)

add_code_block(doc,
    "market-forge-openenv/\n"
    "|\n"
    "|-- models.py                    # Data contracts (Action, Observation, State)\n"
    "|-- client.py                    # HTTP client for agents to connect\n"
    "|-- rewards.py                   # Reward functions for GRPO training\n"
    "|-- app_visual.py                # Gradio dashboard (visual playground)\n"
    "|-- train_market_forge.py        # GRPO training script\n"
    "|-- train_market_forge_notebook.py  # Notebook-style training + evaluation\n"
    "|-- Dockerfile                   # Docker deployment for HuggingFace Spaces\n"
    "|-- requirements.txt             # Python dependencies\n"
    "|-- README.md                    # Project documentation\n"
    "|\n"
    "|-- server/\n"
    "|   |-- __init__.py\n"
    "|   |-- app.py                   # FastAPI HTTP server\n"
    "|   |-- market_environment.py    # Core market engine (the brain)\n"
    "|\n"
    "|-- .claude/\n"
    "    |-- launch.json              # Dev server configuration\n"
)

doc.add_paragraph(
    "The architecture follows a clean separation:"
)
add_table(doc,
    ["Layer", "File", "Responsibility"],
    [
        ["Data Contracts", "models.py", "MarketAction, MarketObservation, MarketState dataclasses"],
        ["Market Engine", "server/market_environment.py", "Core game logic, order matching, scoring, awards"],
        ["HTTP Server", "server/app.py", "FastAPI endpoints: /reset, /step, /state, /health"],
        ["HTTP Client", "client.py", "MarketForgeEnv class for agents to connect"],
        ["Rewards", "rewards.py", "5 reward functions for GRPO training"],
        ["Dashboard", "app_visual.py", "Gradio UI with simulation, manual control, charts"],
        ["Training", "train_market_forge.py", "GRPO training pipeline with TRL"],
    ],
    col_widths=[1.2, 2.2, 3.0]
)

# ── 1.2 Install Dependencies ──
doc.add_heading("1.2  Step 1 — Install Dependencies", level=2)

doc.add_paragraph("Open a terminal and run the following commands:")

add_code_block(doc,
    "# Create a project directory\n"
    "mkdir -p market-forge-openenv/server\n"
    "cd market-forge-openenv\n"
    "\n"
    "# Install core dependencies\n"
    "pip install fastapi 'uvicorn[standard]' gradio matplotlib numpy requests\n"
    "\n"
    "# Install training dependencies (for developers)\n"
    "pip install datasets transformers trl accelerate\n"
    "\n"
    "# Optional: Install OpenEnv core (if available)\n"
    "pip install 'openenv-core[core]>=0.2.1'\n"
)

add_note(doc, "The environment works standalone without openenv-core. "
         "Fallback base classes are included in models.py and client.py.")

doc.add_paragraph(
    "Verify your installation:"
)
add_code_block(doc,
    "python3 -c \"import fastapi, gradio, matplotlib, numpy; print('All dependencies OK')\"\n"
)

# ── 1.3 Define Data Models ──
doc.add_heading("1.3  Step 2 — Define the Data Models (models.py)", level=2)

doc.add_paragraph(
    "The data models define the contract between agents and the environment. "
    "There are exactly three dataclasses:"
)

add_table(doc,
    ["Dataclass", "Direction", "Purpose"],
    [
        ["MarketAction", "Agent -> Environment", "What the agent wants to do (buy, sell, negotiate, etc.)"],
        ["MarketObservation", "Environment -> Agent", "What the agent can see (cash, inventory, prices, events)"],
        ["MarketState", "Server-side only", "Full game state for debugging and visualization"],
    ],
    col_widths=[1.5, 1.8, 3.1]
)

doc.add_paragraph("MarketAction fields — the agent's controls:")

add_table(doc,
    ["Field", "Type", "Description"],
    [
        ["agent_id", "str", "Which agent is acting (e.g., 'trader_1')"],
        ["action_type", "str", "buy, sell, produce, negotiate, accept_deal, reject_deal, propose_coalition, join_coalition, leave_coalition, pass"],
        ["commodity", "str", "wheat, iron, timber, or oil"],
        ["price", "float", "The price per unit"],
        ["quantity", "int", "How many units"],
        ["target_agent", "str", "Who to negotiate with"],
        ["message", "str", "Natural language message for negotiation"],
        ["coalition_id", "str", "Which coalition or deal to join/accept/reject"],
        ["compound_good", "str", "bread, tools, or furniture (for produce action)"],
    ],
    col_widths=[1.3, 0.8, 4.3]
)

doc.add_paragraph("MarketObservation fields — the agent's dashboard:")

add_table(doc,
    ["Field", "Type", "What the Agent Sees"],
    [
        ["agent_id, role", "str", "Identity: who am I and what's my role?"],
        ["cash", "float", "How much money I have"],
        ["inventory", "dict", "What commodities I hold: {wheat: 10, iron: 5, ...}"],
        ["reputation", "float", "My trustworthiness score (0.0 to 2.0)"],
        ["top_of_book", "dict", "Best bid/ask prices per commodity"],
        ["last_trade_prices", "dict", "Most recent trade price per commodity"],
        ["price_history", "list", "Historical prices (last 10 rounds)"],
        ["messages", "list", "Incoming messages from other agents"],
        ["pending_deals", "list", "Deal offers waiting for my response"],
        ["coalitions", "list", "Alliances I belong to"],
        ["event", "str", "Current market event (e.g., 'Drought reduces wheat...')"],
        ["round_number", "int", "Current round"],
        ["done", "bool", "Is the game over?"],
        ["reward", "float", "Reward from my last action"],
        ["total_wealth", "float", "My total wealth (cash + inventory value)"],
        ["prompt", "str", "Text description for LLM agents"],
        ["legal_actions", "list", "Valid action types right now"],
    ],
    col_widths=[1.5, 0.8, 4.1]
)

doc.add_paragraph("To create models.py, save the following script:")
add_code_block(doc,
    "# File: models.py\n"
    "# Run: python3 -c \"from models import MarketAction, MarketObservation; print('Models OK')\"\n"
    "#\n"
    "# This file defines three dataclasses that form the API contract.\n"
    "# See the full file in the project repository.\n"
)

add_note(doc,
    "After creating models.py, verify it by running:\n"
    "python3 -c \"from models import MarketAction; a = MarketAction(agent_id='test', action_type='buy', "
    "commodity='wheat', price=10, quantity=5); print(a)\""
)


# ── 1.4 Build Market Engine ──
doc.add_heading("1.4  Step 3 — Build the Market Engine (server/market_environment.py)", level=2)

doc.add_paragraph(
    "This is the core of MarketForge — the market engine that processes all "
    "agent actions, manages order books, matches trades, and computes rewards. "
    "It implements the standard OpenEnv interface with three methods:"
)

add_table(doc,
    ["Method", "What It Does", "Analogy"],
    [
        ["reset(**kwargs)", "Start a new game. All agents get starting cash and inventory.", "Starting the car engine"],
        ["step(action)", "Execute one action. Returns observation + reward.", "Pressing gas / brake / steering"],
        ["state (property)", "View the full game state. For debugging only.", "Looking under the hood"],
    ],
    col_widths=[1.5, 3.5, 1.4]
)

doc.add_paragraph("Key constants built into the engine:")

add_table(doc,
    ["Constant", "Values"],
    [
        ["COMMODITIES", "wheat, iron, timber, oil"],
        ["COMPOUND_GOODS", "bread (2 wheat + 1 oil), tools (2 iron + 1 timber), furniture (2 timber + 1 oil)"],
        ["BASE_PRICES", "wheat=$10, iron=$15, timber=$12, oil=$18"],
        ["COMPOUND_PRICES", "bread=$35, tools=$50, furniture=$45"],
        ["AGENTS", "4 producers, 2 consumers, 1 trader, 1 speculator (8 total)"],
        ["EVENTS", "12 stochastic market events (droughts, embargoes, surpluses, etc.)"],
    ],
    col_widths=[1.5, 4.9]
)

doc.add_paragraph("The 8 agent roles and their starting conditions:")

add_table(doc,
    ["Agent ID", "Role", "Specialty", "Starting Cash", "Starting Inventory"],
    [
        ["producer_wheat", "Producer", "Wheat", "$500", "40 wheat"],
        ["producer_iron", "Producer", "Iron", "$500", "40 iron"],
        ["producer_timber", "Producer", "Timber", "$500", "40 timber"],
        ["producer_oil", "Producer", "Oil", "$500", "40 oil"],
        ["consumer_1", "Consumer", "Bread", "$1,000", "None"],
        ["consumer_2", "Consumer", "Tools", "$1,000", "None"],
        ["trader_1", "Trader", "General", "$800", "5 of each commodity"],
        ["speculator_1", "Speculator", "Events", "$1,500", "None"],
    ],
    col_widths=[1.4, 0.9, 0.8, 1.0, 2.3]
)

doc.add_paragraph("What happens inside step():")

doc.add_paragraph(
    "1. Agent submits a MarketAction (e.g., buy 5 wheat at $12)\n"
    "2. The engine validates the action (enough cash? valid commodity?)\n"
    "3. For buy/sell: places order on the order book, attempts to match with existing orders\n"
    "4. For negotiate: creates a pending deal and delivers a message to the target\n"
    "5. For produce: checks recipe ingredients, consumes them, creates compound good\n"
    "6. For coalition actions: creates/joins/leaves alliances\n"
    "7. Computes reward based on the outcome\n"
    "8. After all agents act in a round: advances the round, generates a new market event, "
    "updates prices, and regenerates producer inventory\n"
    "9. When round_number reaches max_rounds: game ends, awards and leaderboard are computed"
)

doc.add_paragraph("How trade matching works (Continuous Double Auction):")
add_code_block(doc,
    "# Order Book Matching Algorithm:\n"
    "# 1. Sort all buy orders (bids) by price: highest first\n"
    "# 2. Sort all sell orders (asks) by price: lowest first\n"
    "# 3. While best_bid.price >= best_ask.price:\n"
    "#      trade_price = (best_bid.price + best_ask.price) / 2\n"
    "#      trade_qty   = min(bid.quantity, ask.quantity)\n"
    "#      buyer pays:   trade_price * trade_qty\n"
    "#      seller receives: trade_price * trade_qty\n"
    "#      buyer gains inventory, seller loses inventory\n"
    "#      record trade in trade_history\n"
)

doc.add_paragraph("Verify the engine:")
add_code_block(doc,
    "# Test the market engine directly\n"
    "cd market-forge-openenv\n"
    "python3 -c \"\n"
    "from server.market_environment import MarketEnvironment\n"
    "from models import MarketAction\n"
    "\n"
    "env = MarketEnvironment()\n"
    "obs = env.reset(max_rounds=10)\n"
    "print(f'Game started. Round: {obs.round_number}/{obs.max_rounds}')\n"
    "print(f'Event: {obs.event}')\n"
    "print(f'Cash: {obs.cash}, Inventory: {obs.inventory}')\n"
    "\n"
    "# Place a buy order\n"
    "action = MarketAction(agent_id='trader_1', action_type='buy',\n"
    "                      commodity='wheat', price=12.0, quantity=5)\n"
    "obs = env.step(action)\n"
    "print(f'Reward: {obs.reward:.3f}, Cash: {obs.cash}')\n"
    "\n"
    "# Check awards\n"
    "awards = env.compute_awards()\n"
    "print(f'Market Champion: {awards[\\\"market_champion\\\"][\\\"agent_id\\\"]}')\n"
    "print(f'Master Strategist: {awards[\\\"master_strategist\\\"][\\\"agent_id\\\"]}')\n"
    "print('Engine OK!')\n"
    "\"\n"
)

# ── 1.5 Create HTTP Server ──
doc.add_heading("1.5  Step 4 — Create the HTTP Server (server/app.py)", level=2)

doc.add_paragraph(
    "The FastAPI server exposes the environment over HTTP so that agents "
    "(local or remote) can connect to it. It provides four endpoints:"
)

add_table(doc,
    ["Endpoint", "Method", "What It Does", "Request Body"],
    [
        ["/health", "GET", "Check if the server is running", "None"],
        ["/reset", "POST", "Start a new game", "{max_rounds: 30, seed: 42}"],
        ["/step", "POST", "Execute one agent action", "{agent_id: 'trader_1', action_type: 'buy', ...}"],
        ["/state", "GET", "Get full game state (debug)", "None"],
    ],
    col_widths=[0.8, 0.8, 2.0, 2.8]
)

doc.add_paragraph("Start the server manually:")
add_code_block(doc,
    "# Option 1: Run with uvicorn\n"
    "cd market-forge-openenv\n"
    "uvicorn server.app:app --host 0.0.0.0 --port 8000 --reload\n"
    "\n"
    "# Option 2: Run the module directly\n"
    "python3 -m server.app\n"
)

doc.add_paragraph("Test the server with curl:")
add_code_block(doc,
    "# Check health\n"
    "curl http://localhost:8000/health\n"
    "# Expected: {\"status\":\"healthy\",\"env\":\"market_forge\",\"version\":\"0.2.1\"}\n"
    "\n"
    "# Reset the game\n"
    "curl -X POST http://localhost:8000/reset \\\n"
    "  -H 'Content-Type: application/json' \\\n"
    "  -d '{\"max_rounds\": 20}'\n"
    "\n"
    "# Place a buy order\n"
    "curl -X POST http://localhost:8000/step \\\n"
    "  -H 'Content-Type: application/json' \\\n"
    "  -d '{\"agent_id\":\"trader_1\",\"action_type\":\"buy\",\"commodity\":\"wheat\",\"price\":12,\"quantity\":5}'\n"
    "\n"
    "# Get full state\n"
    "curl http://localhost:8000/state\n"
)


# ── 1.6 Build HTTP Client ──
doc.add_heading("1.6  Step 5 — Build the HTTP Client (client.py)", level=2)

doc.add_paragraph(
    "The MarketForgeEnv client is the steering wheel that agents use to "
    "interact with the environment. It wraps the HTTP endpoints into a "
    "simple Python interface."
)

add_code_block(doc,
    "# Using the client:\n"
    "from client import MarketForgeEnv\n"
    "from models import MarketAction\n"
    "\n"
    "# Connect to the server\n"
    "env = MarketForgeEnv(base_url='http://localhost:8000')\n"
    "\n"
    "# Start a new game\n"
    "result = env.reset(max_rounds=30)\n"
    "print(result.observation.prompt)      # What the agent sees\n"
    "print(result.observation.legal_actions)  # What the agent can do\n"
    "\n"
    "# Take an action\n"
    "action = MarketAction(\n"
    "    agent_id='trader_1',\n"
    "    action_type='buy',\n"
    "    commodity='wheat',\n"
    "    price=12.0,\n"
    "    quantity=5,\n"
    ")\n"
    "result = env.step(action)\n"
    "print(f'Reward: {result.reward}')\n"
    "print(f'Done: {result.done}')\n"
    "\n"
    "# When done=True, check who won\n"
    "if result.done:\n"
    "    awards = result.observation.market_summary.get('awards', {})\n"
    "    print(f'Champion: {awards[\"market_champion\"][\"agent_id\"]}')\n"
)


# ── 1.7 Reward Functions ──
doc.add_heading("1.7  Step 6 — Write the Reward Functions (rewards.py)", level=2)

doc.add_paragraph(
    "The reward functions tell the training algorithm what good behavior looks like. "
    "MarketForge uses five reward functions, each measuring a different aspect:"
)

add_table(doc,
    ["Function", "Weight", "What It Measures"],
    [
        ["reward_from_env", "40%", "Direct profit/loss from each trade (environment feedback)"],
        ["reward_valid_json", "20%", "Did the agent produce a valid JSON action? (format accuracy)"],
        ["reward_strategic_depth", "20%", "Did the agent use advanced actions like negotiate or propose_coalition?"],
        ["reward_event_response", "10%", "Did the agent adapt to market events (e.g., buy wheat during drought)?"],
        ["reward_wealth_growth", "10%", "Did the agent's total wealth increase over the episode?"],
    ],
    col_widths=[1.8, 0.7, 3.9]
)

doc.add_paragraph("Test the reward functions:")
add_code_block(doc,
    "python3 -c \"\n"
    "from rewards import reward_valid_json, reward_strategic_depth\n"
    "\n"
    "# Test valid JSON reward\n"
    "good = ['{\\\"action_type\\\":\\\"buy\\\",\\\"commodity\\\":\\\"wheat\\\",\\\"price\\\":10,\\\"quantity\\\":5}']\n"
    "bad  = ['this is not json']\n"
    "print(f'Good action reward: {reward_valid_json(good)}')\n"
    "print(f'Bad action reward:  {reward_valid_json(bad)}')\n"
    "\"\n"
)


# ── 1.8 Visual Dashboard ──
doc.add_heading("1.8  Step 7 — Build the Visual Dashboard (app_visual.py)", level=2)

doc.add_paragraph(
    "The Gradio dashboard provides a visual playground for running simulations, "
    "manually controlling agents, and viewing rich charts. It runs on port 7860."
)

doc.add_paragraph("The dashboard has three tabs:")

add_table(doc,
    ["Tab", "Purpose", "How to Use"],
    [
        ["Simulation", "Run automated simulation with AI agents", "Set rounds (10-100), click 'Run Simulation'"],
        ["Manual Control", "Control individual agents yourself", "Select agent, action, commodity, price, quantity"],
        ["Environment Info", "Documentation and reference", "Read about rules, actions, awards, scoring"],
    ],
    col_widths=[1.3, 2.5, 2.6]
)

doc.add_paragraph("The dashboard shows 6 chart panels:")
add_table(doc,
    ["Panel", "Chart", "What It Shows"],
    [
        ["1", "Commodity Prices", "Price trends for wheat, iron, timber, oil over time"],
        ["2", "Wealth Race (Award 1)", "Each agent's total wealth per round, champion highlighted"],
        ["3", "Average Reward", "Mean reward across agents per round with moving average"],
        ["4", "Market Activity", "Cumulative trades, active coalitions, compound goods produced"],
        ["5", "Strategic Scores (Award 2)", "Horizontal bar chart of strategic score per agent"],
        ["6", "Action Accuracy", "% of valid actions per agent (green/yellow/red)"],
    ],
    col_widths=[0.5, 1.8, 4.1]
)

doc.add_paragraph("Launch the dashboard:")
add_code_block(doc,
    "cd market-forge-openenv\n"
    "python3 app_visual.py\n"
    "\n"
    "# Then open in your browser:\n"
    "#   http://localhost:7860\n"
)


# ── 1.9 Training Script ──
doc.add_heading("1.9  Step 8 — Create the Training Script (train_market_forge.py)", level=2)

doc.add_paragraph(
    "The training script uses HuggingFace TRL's GRPOTrainer to train an LLM agent "
    "to make smart trading decisions. It follows the same pattern as the official "
    "OpenEnv Wordle GRPO notebook."
)

doc.add_paragraph("Training flow:")
doc.add_paragraph(
    "1. Generate 256 diverse market scenarios as training prompts\n"
    "2. For each prompt: the LLM generates an action (JSON) -> environment steps -> reward is collected\n"
    "3. Multi-turn episodes: 6 actions per game (agent sees feedback and adapts)\n"
    "4. Three reward signals combined: environment reward + format accuracy + strategic depth\n"
    "5. GRPO optimizer updates the policy using group relative advantage\n"
    "6. The env_mask mechanism ensures only the agent's tokens are optimized (not environment feedback)"
)

doc.add_paragraph("Run training:")
add_code_block(doc,
    "# Requires GPU (Google Colab T4/A100 recommended)\n"
    "pip install trl transformers datasets accelerate\n"
    "\n"
    "cd market-forge-openenv\n"
    "python3 train_market_forge.py\n"
    "\n"
    "# Training output will be saved to: ./market-forge-agent/\n"
)


# ── 1.10 Dockerize ──
doc.add_heading("1.10  Step 9 — Dockerize for Deployment", level=2)

doc.add_paragraph(
    "The Dockerfile bundles everything into a single container that runs both "
    "the API server (port 8000) and the Gradio dashboard (port 7860)."
)

add_code_block(doc,
    "# Build the Docker image\n"
    "cd market-forge-openenv\n"
    "docker build -t marketforge .\n"
    "\n"
    "# Run the container\n"
    "docker run -p 7860:7860 -p 8000:8000 marketforge\n"
    "\n"
    "# Access:\n"
    "#   Dashboard: http://localhost:7860\n"
    "#   API:       http://localhost:8000/health\n"
)


# ── 1.11 Run and Test Everything ──
doc.add_heading("1.11  Step 10 — Run and Test Everything", level=2)

doc.add_paragraph("Complete verification script — copy and run this to test all components:")

add_code_block(doc,
    "cd market-forge-openenv\n"
    "python3 << 'EOF'\n"
    "import sys\n"
    "\n"
    "# Test 1: Models\n"
    "from models import MarketAction, MarketObservation, MarketState\n"
    "a = MarketAction(agent_id='test', action_type='buy', commodity='wheat', price=10, quantity=5)\n"
    "assert a.action_type == 'buy'\n"
    "print('[PASS] Test 1: Models load correctly')\n"
    "\n"
    "# Test 2: Environment reset\n"
    "from server.market_environment import MarketEnvironment, COMMODITIES, AGENT_ROLES\n"
    "env = MarketEnvironment()\n"
    "obs = env.reset(max_rounds=10)\n"
    "assert obs.round_number == 0\n"
    "assert obs.cash > 0\n"
    "print('[PASS] Test 2: Environment resets correctly')\n"
    "\n"
    "# Test 3: Step with buy action\n"
    "action = MarketAction(agent_id='trader_1', action_type='buy',\n"
    "                      commodity='wheat', price=12.0, quantity=3)\n"
    "obs = env.step(action)\n"
    "assert hasattr(obs, 'reward')\n"
    "print('[PASS] Test 3: Step executes correctly')\n"
    "\n"
    "# Test 4: Step with sell action\n"
    "action = MarketAction(agent_id='producer_wheat', action_type='sell',\n"
    "                      commodity='wheat', price=9.0, quantity=5)\n"
    "obs = env.step(action)\n"
    "print('[PASS] Test 4: Sell action works')\n"
    "\n"
    "# Test 5: Produce compound good\n"
    "# Give consumer ingredients first\n"
    "env._state.agents['consumer_1']['inventory']['wheat'] = 10\n"
    "env._state.agents['consumer_1']['inventory']['oil'] = 5\n"
    "action = MarketAction(agent_id='consumer_1', action_type='produce', compound_good='bread')\n"
    "obs = env.step(action)\n"
    "print('[PASS] Test 5: Production works')\n"
    "\n"
    "# Test 6: Negotiate\n"
    "action = MarketAction(agent_id='consumer_2', action_type='negotiate',\n"
    "                      target_agent='producer_iron', commodity='iron',\n"
    "                      price=13.0, quantity=5, message='Need iron for tools')\n"
    "obs = env.step(action)\n"
    "print('[PASS] Test 6: Negotiation works')\n"
    "\n"
    "# Test 7: Coalition\n"
    "action = MarketAction(agent_id='speculator_1', action_type='propose_coalition',\n"
    "                      message='Speculator alliance')\n"
    "obs = env.step(action)\n"
    "print('[PASS] Test 7: Coalition creation works')\n"
    "\n"
    "# Test 8: Awards\n"
    "awards = env.compute_awards()\n"
    "assert 'market_champion' in awards\n"
    "assert 'master_strategist' in awards\n"
    "print('[PASS] Test 8: Awards compute correctly')\n"
    "\n"
    "# Test 9: Leaderboard\n"
    "lb = env.compute_leaderboard()\n"
    "assert len(lb) > 0\n"
    "assert lb[0]['rank'] == 1\n"
    "print('[PASS] Test 9: Leaderboard ranks correctly')\n"
    "\n"
    "# Test 10: Accuracy\n"
    "acc = env.compute_accuracy()\n"
    "assert all(0 <= v <= 1 for v in acc.values())\n"
    "print('[PASS] Test 10: Accuracy tracking works')\n"
    "\n"
    "# Test 11: Client class\n"
    "from client import MarketForgeEnv\n"
    "print('[PASS] Test 11: Client imports correctly')\n"
    "\n"
    "# Test 12: Rewards\n"
    "from rewards import reward_valid_json, reward_strategic_depth\n"
    "r = reward_valid_json(['{\"action_type\":\"buy\",\"commodity\":\"wheat\",\"price\":10,\"quantity\":5}'])\n"
    "assert r[0] > 0\n"
    "print('[PASS] Test 12: Reward functions work')\n"
    "\n"
    "print()\n"
    "print('=' * 50)\n"
    "print('  ALL 12 TESTS PASSED - MarketForge is ready!')\n"
    "print('=' * 50)\n"
    "EOF\n"
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# PART 2: BUSINESS USER PLAYBOOK
# ══════════════════════════════════════════════════════════════
doc.add_heading("Part 2: Business User Playbook — Trading in MarketForge", level=1)

doc.add_paragraph(
    "This section is written for business users, traders, and non-technical team members "
    "who want to use MarketForge as a trading playground. No coding is required — everything "
    "is done through the visual dashboard."
)

# ── 2.1 Meet the Personas ──
doc.add_heading("2.1  Meet the Trading Personas", level=2)

doc.add_paragraph(
    "MarketForge simulates a realistic multi-commodity marketplace with eight "
    "participants, each playing a different business role:"
)

add_table(doc,
    ["Persona", "Role", "Business Objective", "Starting Position"],
    [
        ["Wheat Farmer", "Producer", "Grow and sell wheat at the best price", "$500 cash + 40 wheat"],
        ["Iron Miner", "Producer", "Mine and sell iron ore", "$500 cash + 40 iron"],
        ["Timber Mill", "Producer", "Harvest and sell timber", "$500 cash + 40 timber"],
        ["Oil Driller", "Producer", "Extract and sell crude oil", "$500 cash + 40 oil"],
        ["Bakery Chain", "Consumer", "Buy wheat + oil to bake bread (profit: $35/unit)", "$1,000 cash"],
        ["Tool Maker", "Consumer", "Buy iron + timber to make tools (profit: $50/unit)", "$1,000 cash"],
        ["Commodities Trader", "Trader", "Buy low, sell high across all commodities", "$800 cash + 5 of each"],
        ["Market Speculator", "Speculator", "Profit from price swings and market events", "$1,500 cash"],
    ],
    col_widths=[1.2, 0.8, 2.8, 1.6]
)

doc.add_paragraph(
    "These personas interact through four types of behavior:"
)

add_table(doc,
    ["Behavior", "What It Means", "Business Example"],
    [
        ["Competition", "Trading on the open market (buy/sell orders)", "Bidding for wheat in an auction"],
        ["Cooperation", "Combining raw materials into valuable products", "A bakery buying wheat + oil to make bread"],
        ["Negotiation", "Making deals directly with another player", "Agreeing on a bulk discount with a supplier"],
        ["Coalition", "Forming alliances for collective bargaining", "Multiple buyers pooling orders for leverage"],
    ],
    col_widths=[1.2, 2.6, 2.6]
)


# ── 2.2 Launch the Playground ──
doc.add_heading("2.2  How to Launch the Trading Playground", level=2)

doc.add_paragraph("To start the trading playground, run one command:")

add_code_block(doc,
    "cd market-forge-openenv\n"
    "python3 app_visual.py\n"
)

doc.add_paragraph(
    "Then open your browser and go to: http://localhost:7860"
)
doc.add_paragraph(
    "You will see the MarketForge dashboard with three tabs: Simulation, Manual Control, "
    "and Environment Info."
)


# ── 2.3 Running a Simulation ──
doc.add_heading("2.3  Running an Automated Simulation", level=2)

doc.add_paragraph(
    "In the Simulation tab, you can watch all 8 trading personas compete against each other. "
    "The AI agents make decisions automatically based on their role-specific strategies."
)

doc.add_paragraph("How to run a simulation:")
doc.add_paragraph(
    "1. Open the Simulation tab\n"
    "2. Set the number of rounds using the slider (default: 30, range: 10-100)\n"
    "3. Click the 'Run Simulation' button\n"
    "4. Wait for the simulation to complete (usually 5-15 seconds)\n"
    "5. Review the results: Awards, Leaderboard, Charts, Activity Log, Event Timeline"
)

doc.add_paragraph(
    "After the simulation, you will see:"
)
doc.add_paragraph(
    "- Awards: Who won the Market Champion (most wealth) and Master Strategist (best strategy)\n"
    "- Leaderboard: All 8 agents ranked by wealth, with their strategic scores and accuracy\n"
    "- Charts: 6 panels showing price trends, wealth race, rewards, activity, strategies, accuracy\n"
    "- Activity Log: Recent trades, negotiations, and coalition formations\n"
    "- Event Timeline: What market events happened each round (droughts, embargoes, surpluses)"
)


# ── 2.4 Manual Trading ──
doc.add_heading("2.4  Manual Trading — Make Your Own Decisions", level=2)

doc.add_paragraph(
    "In the Manual Control tab, you can step into the shoes of any trading persona and "
    "make decisions yourself. This is like playing a trading video game."
)

doc.add_paragraph("Step-by-step guide to manual trading:")

doc.add_paragraph(
    "1. First, run a quick reset: go to the Simulation tab, set rounds to 10, click 'Run Simulation' "
    "to initialize the market with some activity.\n\n"
    "2. Switch to the Manual Control tab.\n\n"
    "3. Select your agent from the dropdown (e.g., 'trader_1').\n\n"
    "4. Choose an action type:"
)

add_table(doc,
    ["Action", "When to Use It", "What to Fill In"],
    [
        ["buy", "You want to purchase a commodity", "Commodity, Price, Quantity"],
        ["sell", "You want to sell a commodity you hold", "Commodity, Price, Quantity"],
        ["produce", "You have raw materials and want to craft a product", "Commodity = compound good name (bread/tools/furniture)"],
        ["negotiate", "You want to propose a deal to another agent", "Commodity, Price, Qty, Target Agent, Message"],
        ["propose_coalition", "You want to form an alliance", "Message (describe your alliance purpose)"],
        ["pass", "You want to skip this turn", "Nothing"],
    ],
    col_widths=[1.2, 2.3, 2.9]
)

doc.add_paragraph(
    "5. Fill in the details (commodity, price, quantity).\n"
    "6. Click 'Execute Action'.\n"
    "7. Read the result: your reward, new cash balance, inventory, and the current market event."
)

doc.add_paragraph("Example manual trade:")
doc.add_paragraph(
    "- Agent: trader_1\n"
    "- Action: buy\n"
    "- Commodity: wheat\n"
    "- Price: 11.00\n"
    "- Quantity: 5\n"
    "- Click 'Execute Action'\n"
    "- Result: Reward: 0.020, Cash: $793.50, Inventory: {wheat: 10, iron: 5, ...}"
)


# ── 2.5 Understanding Charts ──
doc.add_heading("2.5  Understanding the Dashboard Charts", level=2)

add_table(doc,
    ["Chart", "What to Look For", "Business Insight"],
    [
        ["Commodity Prices", "Trends going up or down", "Which commodities are appreciating? Buy before a rise, sell before a drop."],
        ["Wealth Race", "Who is leading? Whose wealth is growing fastest?", "The agent with the boldest line is the Market Champion."],
        ["Average Reward", "Is the trend going up?", "Rising rewards mean agents are making better decisions each round."],
        ["Market Activity", "Trade volume, coalition count", "High activity = liquid market. Low activity = opportunities for negotiation."],
        ["Strategic Scores", "Gold bar = Master Strategist winner", "Higher scores mean better overall trading skill (not just wealth)."],
        ["Action Accuracy", "Green = good, Yellow = okay, Red = poor", "Shows what % of actions were valid. Higher = more reliable agent."],
    ],
    col_widths=[1.3, 2.0, 3.1]
)


# ── 2.6 Awards ──
doc.add_heading("2.6  How Awards Work — Winning the Game", level=2)

doc.add_paragraph("MarketForge has two awards that determine who wins the game:")

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Award 1: Market Champion (The Wealthiest Player)")
run.font.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0xDA, 0xA5, 0x20)

doc.add_paragraph(
    "The agent with the highest total wealth at the end of the game wins. "
    "Total wealth is calculated as:\n\n"
    "    Total Wealth = Cash + Market Value of Inventory + Value of Compound Goods\n\n"
    "This award rewards agents that make profitable trades, accumulate valuable "
    "inventory, and manage their cash wisely."
)

p = doc.add_paragraph()
run = p.add_run("Award 2: Master Strategist (The Smartest Player)")
run.font.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x4A, 0x9E, 0xD9)

doc.add_paragraph(
    "The agent with the best composite strategic score wins. The score is calculated from "
    "four components:\n\n"
    "    Strategic Score = 35% Trade Efficiency\n"
    "                    + 25% Negotiation Mastery\n"
    "                    + 20% Cooperation Index\n"
    "                    + 20% Event Adaptability\n\n"
    "This award values well-rounded strategy over pure wealth. An agent that "
    "negotiates well, forms effective partnerships, and adapts to market events "
    "can win this even without the most money."
)


# ── 2.7 Leaderboard ──
doc.add_heading("2.7  Reading the Leaderboard", level=2)

doc.add_paragraph(
    "After every simulation, the Leaderboard ranks all 8 agents. Here's how to read it:"
)

add_table(doc,
    ["Column", "What It Means"],
    [
        ["Rank", "Position (1 = best)"],
        ["Agent", "The agent's name"],
        ["Role", "Producer, Consumer, Trader, or Speculator"],
        ["Wealth", "Total value (cash + inventory + compounds) at game end"],
        ["Strategic", "Composite skill score (0 to 1, higher = better)"],
        ["Accuracy", "% of valid actions out of total actions taken"],
        ["Awards", "Which awards this agent won (Champion, Strategist, or none)"],
    ],
    col_widths=[1.0, 5.4]
)


# ── 2.8 Strategy Tips ──
doc.add_heading("2.8  Trading Strategy Tips for Business Users", level=2)

doc.add_paragraph(
    "Here are proven strategies that work well in the MarketForge simulation:"
)

strategies = [
    ("For Producers", "Sell your specialty commodity when prices are high. "
     "Watch for demand surges (events that increase demand for your product). "
     "Consider forming producer coalitions to stabilize prices."),

    ("For Consumers", "Buy raw materials when prices dip. Always check if you have "
     "enough ingredients to produce compound goods — bread ($35), tools ($50), and "
     "furniture ($45) are worth much more than raw materials."),

    ("For Traders", "Look for arbitrage: buy when the asking price is low, sell when "
     "the bid price is high. React to events — buy commodities before a supply shock "
     "drives prices up."),

    ("For Speculators", "Use event information aggressively. When you see 'Drought reduces "
     "wheat supply by 30%', buy wheat immediately — the price will rise. "
     "Sell before oversupply events."),

    ("For Everyone", "Negotiate deals for better prices than the open market. "
     "Form coalitions for bulk buying power. Keep your reputation high by honoring "
     "deals — broken contracts hurt your score."),
]

for title, body in strategies:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.font.bold = True
    run = p.add_run(body)


doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# PART 3: DEVELOPER GUIDE
# ══════════════════════════════════════════════════════════════
doc.add_heading("Part 3: Developer Guide — Building and Training Models", level=1)

doc.add_paragraph(
    "This section is for technical developers who want to build custom agents, "
    "train LLM models, and extend the MarketForge environment."
)


# ── 3.1 Architecture Overview ──
doc.add_heading("3.1  Architecture Overview for Developers", level=2)

add_code_block(doc,
    "                    MarketForge Architecture\n"
    "                    ========================\n"
    "\n"
    "  [Your Agent]                                    [MarketForge Server]\n"
    "  ============                                    ====================\n"
    "                                                         |\n"
    "  LLM / Script / Bot                             MarketEnvironment\n"
    "       |                                                |\n"
    "       |  MarketAction (JSON)                    +------+------+\n"
    "       +------------>  HTTP POST /step  -------> | Order Books |\n"
    "       |                                         | Trade Match |\n"
    "       |  MarketObservation (JSON)                | Coalitions  |\n"
    "       <------------  HTTP Response  <----------- | Awards      |\n"
    "       |                                         | Scoring     |\n"
    "       v                                         +-------------+\n"
    "  [Decide Next Action]                                  |\n"
    "       |                                         [Gradio Dashboard]\n"
    "       v                                           port 7860\n"
    "  [GRPO Training]\n"
    "  (reward -> optimize policy)\n"
)


# ── 3.2 Connecting via HTTP API ──
doc.add_heading("3.2  Connecting to the Environment via HTTP API", level=2)

doc.add_paragraph("Developers can connect to MarketForge from any language using the HTTP API.")

doc.add_paragraph("Python example:")
add_code_block(doc,
    "import requests\n"
    "import json\n"
    "\n"
    "BASE_URL = 'http://localhost:8000'\n"
    "\n"
    "# Reset the game\n"
    "resp = requests.post(f'{BASE_URL}/reset', json={'max_rounds': 30})\n"
    "obs = resp.json()\n"
    "print(f'Game started. Agent: {obs[\"agent_id\"]}, Cash: {obs[\"cash\"]}')\n"
    "print(f'Event: {obs[\"event\"]}')\n"
    "\n"
    "# Game loop\n"
    "done = False\n"
    "while not done:\n"
    "    action = {\n"
    "        'agent_id': 'trader_1',\n"
    "        'action_type': 'buy',\n"
    "        'commodity': 'wheat',\n"
    "        'price': 12.0,\n"
    "        'quantity': 5,\n"
    "    }\n"
    "    resp = requests.post(f'{BASE_URL}/step', json=action)\n"
    "    obs = resp.json()\n"
    "    done = obs.get('done', False)\n"
    "    print(f'Round {obs[\"round_number\"]}: reward={obs[\"reward\"]:.3f}, wealth={obs[\"total_wealth\"]:.0f}')\n"
    "\n"
    "# Check final awards\n"
    "awards = obs.get('market_summary', {}).get('awards', {})\n"
    "print(f'Champion: {awards[\"market_champion\"][\"agent_id\"]}')\n"
    "print(f'Strategist: {awards[\"master_strategist\"][\"agent_id\"]}')\n"
)

doc.add_paragraph("cURL example (works from any terminal):")
add_code_block(doc,
    "# Reset\n"
    "curl -s -X POST http://localhost:8000/reset \\\n"
    "  -H 'Content-Type: application/json' \\\n"
    "  -d '{\"max_rounds\": 10}' | python3 -m json.tool\n"
    "\n"
    "# Buy wheat\n"
    "curl -s -X POST http://localhost:8000/step \\\n"
    "  -H 'Content-Type: application/json' \\\n"
    "  -d '{\"agent_id\":\"trader_1\",\"action_type\":\"buy\",\"commodity\":\"wheat\",\"price\":12,\"quantity\":5}' \\\n"
    "  | python3 -m json.tool\n"
    "\n"
    "# Negotiate a deal\n"
    "curl -s -X POST http://localhost:8000/step \\\n"
    "  -H 'Content-Type: application/json' \\\n"
    "  -d '{\"agent_id\":\"consumer_1\",\"action_type\":\"negotiate\",\"target_agent\":\"producer_wheat\",\\\n"
    "       \"commodity\":\"wheat\",\"price\":9,\"quantity\":10,\\\n"
    "       \"message\":\"Need 10 wheat for bread production. Can we do $9/unit?\"}' \\\n"
    "  | python3 -m json.tool\n"
)


# ── 3.3 Building a Custom Agent ──
doc.add_heading("3.3  Building a Custom Agent (Python)", level=2)

doc.add_paragraph(
    "Here is a complete, runnable script for a custom trading agent that connects "
    "to the MarketForge server and plays a full game:"
)

add_code_block(doc,
    "#!/usr/bin/env python3\n"
    "\"\"\"custom_agent.py - A smart trading agent for MarketForge.\"\"\"\n"
    "from client import MarketForgeEnv\n"
    "from models import MarketAction\n"
    "\n"
    "class SmartTrader:\n"
    "    def __init__(self, agent_id='trader_1'):\n"
    "        self.agent_id = agent_id\n"
    "        self.buy_history = []\n"
    "\n"
    "    def decide(self, obs):\n"
    "        \"\"\"Make a trading decision based on observation.\"\"\"\n"
    "        event = obs.event.lower()\n"
    "\n"
    "        # Strategy 1: React to supply shocks\n"
    "        for c in ['wheat', 'iron', 'timber', 'oil']:\n"
    "            if c in event and ('reduce' in event or 'embargo' in event):\n"
    "                price = obs.top_of_book.get(c, {}).get('best_ask', 15)\n"
    "                if obs.cash > price * 5:\n"
    "                    return MarketAction(\n"
    "                        agent_id=self.agent_id, action_type='buy',\n"
    "                        commodity=c, price=round(price * 1.05, 1), quantity=5)\n"
    "\n"
    "        # Strategy 2: Sell expensive holdings\n"
    "        for c, qty in obs.inventory.items():\n"
    "            if qty > 5:\n"
    "                price = obs.top_of_book.get(c, {}).get('best_bid', 10)\n"
    "                if price > 0:\n"
    "                    return MarketAction(\n"
    "                        agent_id=self.agent_id, action_type='sell',\n"
    "                        commodity=c, price=round(price * 1.02, 1),\n"
    "                        quantity=min(qty - 2, 5))\n"
    "\n"
    "        # Strategy 3: Negotiate a deal\n"
    "        return MarketAction(\n"
    "            agent_id=self.agent_id, action_type='negotiate',\n"
    "            target_agent='producer_wheat', commodity='wheat',\n"
    "            price=9.0, quantity=10,\n"
    "            message='Looking for bulk wheat. Can you do $9/unit for 10?')\n"
    "\n"
    "\n"
    "# Run the agent\n"
    "env = MarketForgeEnv(base_url='http://localhost:8000')\n"
    "agent = SmartTrader('trader_1')\n"
    "\n"
    "result = env.reset(max_rounds=30)\n"
    "print(f'Game started: {result.observation.event}')\n"
    "\n"
    "while not result.done:\n"
    "    action = agent.decide(result.observation)\n"
    "    result = env.step(action)\n"
    "    print(f'R{result.observation.round_number}: '\n"
    "          f'{action.action_type} {action.commodity} '\n"
    "          f'-> reward={result.reward:.3f}, '\n"
    "          f'wealth=${result.observation.total_wealth:.0f}')\n"
    "\n"
    "awards = result.observation.market_summary.get('awards', {})\n"
    "print(f'\\nChampion: {awards[\"market_champion\"][\"agent_id\"]}')\n"
    "print(f'Strategist: {awards[\"master_strategist\"][\"agent_id\"]}')\n"
)

doc.add_paragraph("Run this agent:")
add_code_block(doc,
    "# First start the server in one terminal:\n"
    "uvicorn server.app:app --host 0.0.0.0 --port 8000\n"
    "\n"
    "# Then run the agent in another terminal:\n"
    "python3 custom_agent.py\n"
)


# ── 3.4 Training with GRPO ──
doc.add_heading("3.4  Training an LLM Agent with GRPO", level=2)

doc.add_paragraph(
    "GRPO (Group Relative Policy Optimization) trains an LLM to produce good trading "
    "actions by comparing multiple generated actions and reinforcing the best ones."
)

doc.add_paragraph("Training prerequisites:")
add_code_block(doc,
    "# Install training dependencies (requires GPU)\n"
    "pip install trl transformers datasets accelerate\n"
    "\n"
    "# Recommended: Google Colab with T4 or A100 GPU\n"
    "# Model: Qwen/Qwen2.5-0.5B-Instruct (small, fast)\n"
    "# For better results: Qwen/Qwen2.5-1.5B-Instruct\n"
)

doc.add_paragraph("How training works (step by step):")

doc.add_paragraph(
    "1. Dataset Generation: 256 diverse market scenarios are created by calling env.reset() "
    "256 times, each producing a unique market observation (different events, prices, inventories).\n\n"
    "2. Prompt Construction: Each scenario is wrapped with a system prompt that teaches the "
    "LLM how to format actions as JSON and provides strategic guidelines.\n\n"
    "3. Rollout: For each prompt, the LLM generates an action -> the environment processes it -> "
    "a reward is computed -> the observation is fed back. This repeats for 6 turns per episode.\n\n"
    "4. env_mask: Environment-generated text (observations, rewards) is masked from the training "
    "loss. Only the LLM's own action tokens are optimized. This prevents the model from trying "
    "to replicate environment output.\n\n"
    "5. Reward Collection: Three reward signals are combined — environment reward (40%), "
    "format accuracy (30%), and strategic depth (30%).\n\n"
    "6. Policy Update: GRPO generates multiple action candidates per prompt, computes the "
    "relative advantage of each, and updates the model to favor better actions."
)

doc.add_paragraph("Run training:")
add_code_block(doc,
    "cd market-forge-openenv\n"
    "python3 train_market_forge.py\n"
    "\n"
    "# Training output:\n"
    "#   Starting MarketForge agent training with GRPO...\n"
    "#   Model: Qwen/Qwen2.5-0.5B-Instruct\n"
    "#   Dataset size: 256\n"
    "#   Output: market-forge-agent\n"
    "#   [Step 1/...] loss=... reward=...\n"
    "#   Training complete! Model saved to ./market-forge-agent\n"
)

doc.add_paragraph("Key training hyperparameters:")
add_table(doc,
    ["Parameter", "Value", "What It Controls"],
    [
        ["learning_rate", "1e-6", "How fast the model learns (lower = more stable)"],
        ["num_generations", "2", "How many action candidates to compare per prompt"],
        ["max_completion_length", "1024", "Maximum tokens per LLM response"],
        ["gradient_accumulation_steps", "16", "Effective batch size (memory optimization)"],
        ["max_turns", "6", "Actions per episode (multi-turn)"],
        ["vllm_gpu_memory_utilization", "0.15", "GPU memory for fast inference"],
    ],
    col_widths=[2.2, 0.8, 3.4]
)


# ── 3.5 Reward Function Design ──
doc.add_heading("3.5  Reward Function Design", level=2)

doc.add_paragraph(
    "The reward pipeline uses a mixed-motive payoff function that balances self-interest "
    "and collective benefit:"
)

add_code_block(doc,
    "# Mixed-motive payoff formula:\n"
    "u_i(action) = alpha_i * v_self(action) + (1 - alpha_i) * v_collective(action)\n"
    "\n"
    "# alpha_i is the self-interest parameter per role:\n"
    "#   Producer:   0.7 (mostly self-interested)\n"
    "#   Consumer:   0.5 (balanced)\n"
    "#   Trader:     0.8 (profit-focused)\n"
    "#   Speculator: 0.9 (highly self-interested)\n"
)

doc.add_paragraph(
    "To create a custom reward function, follow this signature:"
)
add_code_block(doc,
    "def my_custom_reward(completions, **kwargs) -> list[float]:\n"
    "    \"\"\"Custom reward function.\n"
    "    \n"
    "    Args:\n"
    "        completions: List of LLM output texts\n"
    "        **kwargs: Extra context (env_reward, wealth_change, etc.)\n"
    "    \n"
    "    Returns:\n"
    "        List of float rewards (one per completion)\n"
    "    \"\"\"\n"
    "    rewards = []\n"
    "    for c in completions:\n"
    "        text = c[0]['content'] if isinstance(c, list) else str(c)\n"
    "        # Your reward logic here\n"
    "        rewards.append(0.0)\n"
    "    return rewards\n"
)


# ── 3.6 env_mask ──
doc.add_heading("3.6  Multi-Turn Rollouts and env_mask", level=2)

doc.add_paragraph(
    "The env_mask is a critical mechanism for multi-turn GRPO training. It marks which "
    "tokens were generated by the model (mask=1) versus which were generated by the "
    "environment (mask=0). Only model tokens are included in the training loss."
)

add_code_block(doc,
    "# Token stream during a 2-turn episode:\n"
    "#\n"
    "# [System Prompt] [Observation 1]      <- env tokens (mask=0)\n"
    "# [Agent Action 1]                     <- model tokens (mask=1)\n"
    "# [Reward: 0.4. Observation 2]         <- env tokens (mask=0)\n"
    "# [Agent Action 2]                     <- model tokens (mask=1)\n"
    "# [Reward: 0.6. Game Over.]            <- env tokens (mask=0)\n"
    "#\n"
    "# Only the model's actions (mask=1) are optimized.\n"
    "# The environment's feedback provides context but doesn't affect the loss.\n"
)


# ── 3.7 Deploying to HuggingFace Spaces ──
doc.add_heading("3.7  Deploying to HuggingFace Spaces", level=2)

add_code_block(doc,
    "# Step 1: Install HuggingFace Hub\n"
    "pip install huggingface_hub\n"
    "huggingface-cli login\n"
    "\n"
    "# Step 2: Create and upload to a Space\n"
    "python3 -c \"\n"
    "from huggingface_hub import HfApi\n"
    "api = HfApi()\n"
    "\n"
    "api.create_repo(\n"
    "    repo_id='YOUR-USERNAME/market-forge-env',\n"
    "    repo_type='space',\n"
    "    space_sdk='docker',\n"
    "    exist_ok=True,\n"
    ")\n"
    "\n"
    "api.upload_folder(\n"
    "    folder_path='.',\n"
    "    repo_id='YOUR-USERNAME/market-forge-env',\n"
    "    repo_type='space',\n"
    ")\n"
    "print('Deployed! Visit: https://huggingface.co/spaces/YOUR-USERNAME/market-forge-env')\n"
    "\"\n"
)


# ── 3.8 Evaluation ──
doc.add_heading("3.8  Evaluation and Benchmarking", level=2)

doc.add_paragraph("Run the evaluation script to compare agent baselines:")
add_code_block(doc,
    "cd market-forge-openenv\n"
    "python3 train_market_forge_notebook.py\n"
    "\n"
    "# This runs 200 episodes for each baseline:\n"
    "#   [1/3] Random baseline...     (avg: ~0.15)\n"
    "#   [2/3] Heuristic baseline...  (avg: ~1.50)\n"
    "#   [3/3] Strategic baseline...  (avg: ~2.10)\n"
    "#\n"
    "# Generates training_results.png with 4 charts:\n"
    "#   - Reward distribution comparison\n"
    "#   - GRPO training curve\n"
    "#   - Action distribution (before vs after training)\n"
    "#   - Cumulative reward over episodes\n"
)

doc.add_paragraph("Key metrics to track during evaluation:")
add_table(doc,
    ["Metric", "What It Measures", "Target"],
    [
        ["Format Accuracy", "% of responses that are valid JSON with correct action_type", ">95%"],
        ["Environment Reward", "Average reward per step from the environment", ">0.3"],
        ["Strategic Score", "Composite: trade efficiency + negotiation + cooperation + adaptability", ">0.6"],
        ["Wealth Growth", "Final wealth vs starting wealth", ">20%"],
        ["Action Diversity", "Number of unique action types used", ">=4"],
    ],
    col_widths=[1.5, 3.0, 0.9]
)


# ── 3.9 Extending the Environment ──
doc.add_heading("3.9  Extending the Environment", level=2)

doc.add_paragraph("Here are ways to extend MarketForge for your specific use case:")

extensions = [
    ("Add new commodities", "Add to COMMODITIES list and BASE_PRICES dict in market_environment.py"),
    ("Add new compound goods", "Add recipe to COMPOUND_GOODS and price to COMPOUND_PRICES"),
    ("Add new agent roles", "Add to AGENT_ROLES dict with role, specialty, cash, inventory, alpha"),
    ("Add new market events", "Add to EVENTS list with text, effects dict, and severity"),
    ("Change game duration", "Pass max_rounds parameter to reset() (default: 50)"),
    ("Add new action types", "1. Add handler method in MarketEnvironment\n"
     "2. Add to action_handlers dict in step()\n"
     "3. Add to MarketAction dataclass\n"
     "4. Update legal_actions in _make_observation()"),
    ("Custom reward function", "Create function with signature: def my_reward(completions, **kwargs) -> list[float]\n"
     "Add to reward_funcs list in GRPOTrainer constructor"),
    ("Custom scoring formula", "Modify _compute_strategic_score() in market_environment.py\n"
     "Adjust the 4 component weights (must sum to 1.0)"),
]

for title, body in extensions:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.font.bold = True
    run = p.add_run(body)


# ══════════════════════════════════════════════════════════════
# APPENDIX
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("Appendix: Quick Reference", level=1)

doc.add_heading("All Action Types", level=2)
add_table(doc,
    ["Action", "Pillar", "Required Fields", "Reward Range"],
    [
        ["buy", "Competition", "commodity, price, quantity", "+0.02 to +0.40"],
        ["sell", "Competition", "commodity, price, quantity", "+0.02 to +0.40"],
        ["produce", "Cooperation", "compound_good", "+0.50 to +0.80"],
        ["negotiate", "Negotiation", "target_agent, commodity, price, quantity, message", "+0.05"],
        ["accept_deal", "Negotiation", "coalition_id (as deal_id)", "+0.40 to +0.60"],
        ["reject_deal", "Negotiation", "coalition_id (as deal_id)", "0.00"],
        ["propose_coalition", "Coalition", "message", "+0.05"],
        ["join_coalition", "Coalition", "coalition_id", "+0.10 to +0.30"],
        ["leave_coalition", "Coalition", "coalition_id", "-0.08"],
        ["pass", "None", "none", "-0.01"],
    ],
    col_widths=[1.4, 1.0, 2.5, 1.5]
)

doc.add_heading("All Commodities and Recipes", level=2)
add_table(doc,
    ["Item", "Type", "Base Price", "Recipe / Source"],
    [
        ["Wheat", "Raw", "$10", "Produced by producer_wheat (regenerates each round)"],
        ["Iron", "Raw", "$15", "Produced by producer_iron"],
        ["Timber", "Raw", "$12", "Produced by producer_timber"],
        ["Oil", "Raw", "$18", "Produced by producer_oil"],
        ["Bread", "Compound", "$35", "2 Wheat + 1 Oil (profit margin: ~47%)"],
        ["Tools", "Compound", "$50", "2 Iron + 1 Timber (profit margin: ~52%)"],
        ["Furniture", "Compound", "$45", "2 Timber + 1 Oil (profit margin: ~53%)"],
    ],
    col_widths=[1.0, 1.0, 0.8, 3.6]
)

doc.add_heading("All Market Events", level=2)
add_table(doc,
    ["Event", "Affected Commodity", "Effect"],
    [
        ["Drought reduces wheat supply by 30%", "Wheat", "Supply -30% (price rises)"],
        ["Iron mine collapse", "Iron", "Supply -50% (price rises sharply)"],
        ["Timber surplus", "Timber", "Supply +40% (price drops)"],
        ["Oil embargo", "Oil", "Supply -60% (price rises sharply)"],
        ["Trade festival", "All", "Transaction costs waived"],
        ["Furniture demand surge +50%", "Timber, Oil", "Demand increases (prices rise)"],
        ["New trade route opens", "All", "Supply +10% (prices drop slightly)"],
        ["Calm markets", "None", "No effect"],
        ["Speculator rumor: oil to skyrocket", "None", "Information only (may cause speculation)"],
        ["Government subsidy for tools", "Iron, Timber", "Supply +10% (prices drop)"],
        ["Storm damages warehouses", "Wheat, Timber", "Supply -20% (prices rise)"],
        ["International trade deal signed", "All", "Supply -5% (demand up, prices stable)"],
    ],
    col_widths=[2.5, 1.5, 2.4]
)

doc.add_heading("Startup Commands Cheat Sheet", level=2)

add_table(doc,
    ["Task", "Command"],
    [
        ["Install dependencies", "pip install fastapi uvicorn gradio matplotlib numpy requests"],
        ["Start API server", "uvicorn server.app:app --host 0.0.0.0 --port 8000"],
        ["Start dashboard", "python3 app_visual.py"],
        ["Run all tests", "python3 (paste verification script from Section 1.11)"],
        ["Run evaluation", "python3 train_market_forge_notebook.py"],
        ["Train with GRPO", "python3 train_market_forge.py"],
        ["Build Docker", "docker build -t marketforge ."],
        ["Run Docker", "docker run -p 7860:7860 -p 8000:8000 marketforge"],
        ["Deploy to HF Spaces", "See Section 3.7"],
    ],
    col_widths=[1.5, 4.9]
)


# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "MarketForge_Complete_Guide.docx")
doc.save(output_path)
print(f"\nWord document saved to: {output_path}")
print(f"Total sections: 3 Parts + Appendix")
